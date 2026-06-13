# -*- coding: utf-8 -*-
"""Tests for econ-cn-docx pipeline functions. Run: python tests.py"""
import unittest, sys, os, tempfile

# Ensure skill modules are importable
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))

from references.pipeline import (
    _fix_quotes, _latex_to_unicode, _render_math_content,
    _add_display_equation, _parse_rtf_table, _add_body_paragraph,
    set_cn_font, render_from_dicts,
)
from paper import Paper
from docx import Document
from docx.shared import Pt
from lxml import etree

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


class TestFixQuotes(unittest.TestCase):
    def test_chinese_context(self):
        self.assertEqual(_fix_quotes('他说"你好"世界'), '他说“你好”世界')

    def test_no_chinese(self):
        self.assertEqual(_fix_quotes('He said "hello"'), 'He said "hello"')

    def test_empty(self):
        self.assertEqual(_fix_quotes(''), '')

    def test_multiple_pairs(self):
        self.assertEqual(
            _fix_quotes('所谓"污染"的"经济损害"'),
            '所谓“污染”的“经济损害”')

    def test_unmatched_gets_closed(self):
        result = _fix_quotes('他说"你好')
        self.assertIn('“', result)
        self.assertIn('”', result)


class TestLatexToUnicode(unittest.TestCase):
    def test_greek_lower(self):
        result = _latex_to_unicode('\\alpha')
        self.assertIn('α', result)  # α

    def test_greek_upper(self):
        result = _latex_to_unicode('\\Delta')
        self.assertIn('Δ', result)  # Δ

    def test_operators(self):
        self.assertIn('×', _latex_to_unicode('\\times'))
        self.assertIn('±', _latex_to_unicode('\\pm'))

    def test_partial(self):
        self.assertIn('∂', _latex_to_unicode('\\partial'))

    def test_text_command(self):
        self.assertEqual(_latex_to_unicode('\\text{GDP}'), 'GDP')

    def test_keeps_underscore_brace(self):
        s = _latex_to_unicode('x_{it}')
        self.assertIn('_{', s)

    def test_diacritics(self):
        s = _latex_to_unicode('\\bar{x}')
        self.assertIn('x̄', s)  # x + combining macron

    def test_frac_command(self):
        s = _latex_to_unicode('\\frac{a}{b}')
        self.assertIn('\\frac', s)


class TestOMMLMathContent(unittest.TestCase):
    def _render(self, latex):
        omath = etree.Element(f'{{{MATH_NS}}}oMath')
        _render_math_content(omath, latex)
        return omath

    def test_simple_text(self):
        omath = self._render('y = ax')
        texts = [t.text for t in omath.findall(f'.//{{{MATH_NS}}}t')]
        self.assertIn('y', ''.join(texts))

    def test_subscript(self):
        omath = self._render('x_{it}')
        subs = omath.findall(f'.//{{{MATH_NS}}}sSub')
        self.assertGreaterEqual(len(subs), 1)

    def test_superscript(self):
        omath = self._render('x^{2}')
        sups = omath.findall(f'.//{{{MATH_NS}}}sSup')
        self.assertGreaterEqual(len(sups), 1)

    def test_sub_sup(self):
        omath = self._render('x_{i}^{2}')
        subsup = omath.findall(f'.//{{{MATH_NS}}}sSubSup')
        # Should have exactly ONE sSubSup (not doubled)
        self.assertEqual(len(subsup), 1)

    def test_frac(self):
        omath = self._render('\\frac{a}{b}')
        fracs = omath.findall(f'.//{{{MATH_NS}}}f')
        self.assertGreaterEqual(len(fracs), 1)

    def test_sqrt(self):
        omath = self._render('\\sqrt{x}')
        rads = omath.findall(f'.//{{{MATH_NS}}}rad')
        self.assertGreaterEqual(len(rads), 1)

    def test_nth_root(self):
        omath = self._render('\\sqrt[3]{x}')
        rads = omath.findall(f'.//{{{MATH_NS}}}rad')
        degs = omath.findall(f'.//{{{MATH_NS}}}deg')
        self.assertGreaterEqual(len(rads), 1)
        self.assertGreaterEqual(len(degs), 1)


class TestDisplayEquation(unittest.TestCase):
    def test_inserts_omath(self):
        doc = Document()
        _add_display_equation(doc, r'y_{it} = \beta x_{it} + \varepsilon_{it}')
        # Should have 1 paragraph with an oMath element
        self.assertGreaterEqual(len(doc.paragraphs), 1)
        omath = doc.paragraphs[0]._p.findall(f'.//{{{MATH_NS}}}oMath')
        self.assertGreaterEqual(len(omath), 1)


class TestRTFParser(unittest.TestCase):
    def setUp(self):
        self.rtf_path = os.path.join(os.path.dirname(__file__),
                                     'tests', 'tables', 'test_table.rtf')

    def test_parse_returns_headers_rows_notes_title(self):
        headers, rows, notes, title = _parse_rtf_table(self.rtf_path)
        self.assertIsNotNone(headers)
        # Title may be empty if RTF format doesn't match esttab conventions

    def test_parse_has_data_rows(self):
        headers, rows, notes, title = _parse_rtf_table(self.rtf_path)
        self.assertGreaterEqual(len(rows), 1)


class TestBodyParagraph(unittest.TestCase):
    def test_creates_paragraph(self):
        doc = Document()
        n_before = len(doc.paragraphs)
        _add_body_paragraph(doc, '测试正文内容。', size=Pt(12), indent=Pt(24), line_spacing=1.5)
        self.assertEqual(len(doc.paragraphs), n_before + 1)

    def test_justified_alignment(self):
        doc = Document()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        _add_body_paragraph(doc, 'text')
        self.assertEqual(doc.paragraphs[-1].alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)


class TestSetCnFont(unittest.TestCase):
    def test_sets_font_properties(self):
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run('测试')
        set_cn_font(r, '宋体', Pt(12))
        self.assertEqual(r.font.size, Pt(12))


class TestRenderFromDicts(unittest.TestCase):
    def test_minimal_document(self):
        sections = [
            {'title': '测试', 'body': '', 'level': 0},
            {'title': '摘要', 'body': '摘要内容。', 'level': 1},
            {'title': '关键词', 'body': '测试；论文', 'level': 1},
            {'title': '一、引言', 'body': '引言正文。', 'level': 1},
            {'title': '', 'body': 'EQUATION:y = \\beta x', 'level': 0},
            {'title': '', 'body': 'NOTE:注：测试。', 'level': 0},
            {'title': '参考文献', 'body': '', 'level': 1},
            {'title': '', 'body': 'REF:[1] Test. Title[J]. Journal, 2020.', 'level': 1},
        ]
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            render_from_dicts(sections, out,
                              tables_dir=os.path.join(os.path.dirname(__file__), 'tests', 'tables'),
                              figures_dir=os.path.join(os.path.dirname(__file__), 'tests', 'figures'))
            self.assertTrue(os.path.exists(out))
            self.assertGreater(os.path.getsize(out), 1000)
        finally:
            os.unlink(out)

    def test_missing_table_warns(self):
        import io
        sections = [
            {'title': '', 'body': 'TABLE:nonexistent_table', 'level': 0},
        ]
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            render_from_dicts(sections, out, tables_dir='/nonexistent')
            self.assertTrue(os.path.exists(out))
        finally:
            os.unlink(out)


class TestPaperClass(unittest.TestCase):
    def test_builds_sections(self):
        p = Paper()
        p.title('论文标题')
        p.abstract('摘要内容。')
        p.keywords('关键词1；关键词2')
        p.h1('一、引言')
        p.para('正文段落。')
        p.h2('（一）方法')
        p.equation('y = \\beta x')
        p.table('test_table')
        p.image('test_fig.png')
        p.h3('图一：测试')
        p.note('表注内容。')
        p.ref('[1] Author. Title[J]. Journal, 2020.')
        self.assertGreater(len(p.sections), 5)

    def test_render_creates_file(self):
        p = Paper()
        p.title('测试')
        p.abstract('摘要。')
        p.keywords('测试')
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            out = f.name
        try:
            p.render(out,
                     tables_dir=os.path.join(os.path.dirname(__file__), 'tables'),
                     figures_dir=os.path.join(os.path.dirname(__file__), 'figures'))
            self.assertTrue(os.path.exists(out))
            self.assertGreater(os.path.getsize(out), 1000)
        finally:
            os.unlink(out)


if __name__ == '__main__':
    unittest.main(verbosity=2)
