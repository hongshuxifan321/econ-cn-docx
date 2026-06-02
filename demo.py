# -*- coding: utf-8 -*-
'econ-cn-docx 完整演示：三个预设风格 + 所有核心功能'
import tempfile, os
from docx import Document
from cn_docx_toolkit import *

# ═══════════════════════════════════════════════════
# 演示一：毕业论文风格（thesis）
# ═══════════════════════════════════════════════════
doc1 = Document()
setup_page(doc1, style='thesis')

# 封面
add_cover_page(doc1, '空气污染、远程工作可行性与城市经济产出',
                 school='某大学经济学院', author='某同学',
                 date='2026年5月', level='thesis')

# 摘要和关键词
add_abstract_and_keywords(doc1,
    '本文基于2000-2020年中国287个地级市面板数据，'
    '考察远程工作可行性（RWI）是否调节PM2.5污染对城市GDP的损害。',
    'PM2.5；远程工作；城市GDP；面板固定效应')

# 正文
add_numbered_heading(doc1, '引言', level=1)
add_body(doc1, '空气污染损害经济产出（Graff Zivin & Neidell, 2012）。'
    '但一个城市的产业结构决定了它受污染影响的程度——'
    '如果大量从业人员可以在家办公，污染对经济的冲击应当更小。')
add_body(doc1, '本文的边际效应方程为：')
add_display_equation(doc1,
    r'\ln(GDP_{ct}) = \beta_1 PM2.5_{ct} + \beta_2 PM2.5_{ct} \times RWI_c + \gamma \ln(Pop_{ct}) + \alpha_c + \lambda_t + \varepsilon_{ct}')

add_numbered_heading(doc1, '数据来源', level=2)
add_body(doc1, 'PM2.5数据来自CHAP卫星反演数据集（1 km分辨率，2000-2021年），'
    '经济变量来自《中国城市统计年鉴》，RWI基于2020年七普职业分布数据。')

add_numbered_heading(doc1, '实证结果', level=1)
add_numbered_heading(doc1, '基准回归', level=2)
add_body(doc1, '表1报告了基准回归结果。'
    '交互项$PM2.5_{ct} \\times RWI_c$的系数显著为正（$\\beta_2 = +0.0306$, $p < 0.05$），'
    '表明RWI越高的城市，PM2.5对GDP的边际损害越小。')

# 回归表格
models = [
    {'name': '(1) FE', 'N': 5344, 'R2': 0.027,
     'coef': {'PM2.5': -0.0081, 'PM2.5×RWI': 0.0306},
     'se':   {'PM2.5': 0.002, 'PM2.5×RWI': 0.012},
     'pvalues': {'PM2.5': 0.000, 'PM2.5×RWI': 0.012}},
    {'name': '(2) FE', 'N': 4693, 'R2': 0.059,
     'coef': {'PM2.5': -0.0213, 'PM2.5×RWI': 0.1156},
     'se':   {'PM2.5': 0.002, 'PM2.5×RWI': 0.015},
     'pvalues': {'PM2.5': 0.000, 'PM2.5×RWI': 0.000}},
    {'name': '(3) +Wage', 'N': 2510, 'R2': 0.514,
     'coef': {'PM2.5': -0.0182, 'PM2.5×RWI': 0.0333},
     'se':   {'PM2.5': 0.004, 'PM2.5×RWI': 0.025},
     'pvalues': {'PM2.5': 0.000, 'PM2.5×RWI': 0.182}},
]
add_regression_table(doc1, models,
    notes='注：括号内为城市层面聚类标准误。*** p<0.01, ** p<0.05, * p<0.1')

add_numbered_heading(doc1, '稳健性检验', level=2)
add_bullet(doc1, '替换被解释变量为人均GDP：交互项系数+0.1156（$p < 0.001$），通过')
add_bullet(doc1, '排除直辖市：交互项系数+0.0341（$p = 0.014$），通过')
add_bullet(doc1, '提前期检验：Lead PM2.5×RWI不显著（$p = 0.319$），排除反向因果')

# 参考文献
add_heading_1(doc1, '参考文献')
add_reference(doc1, 'Dingel J I, Neiman B. How many jobs can be done at home?[J]. Journal of Public Economics, 2020, 189: 104235.')
add_reference(doc1, 'Graff Zivin J, Neidell M. The impact of pollution on worker productivity[J]. American Economic Review, 2012, 102(7): 3652-3673.')
add_reference(doc1, 'He J, Liu H, Salvo A. Severe air pollution and labor productivity[J]. American Economic Journal: Applied Economics, 2019, 11(1): 173-201.')

add_page_number(doc1, show_total=True)
out1 = os.path.join(tempfile.gettempdir(), 'demo_thesis.docx')
doc1.save(out1)
print(f'演示一（thesis）已保存: {out1}')


# ═══════════════════════════════════════════════════
# 演示二：期刊投稿风格（journal）
# ═══════════════════════════════════════════════════
doc2 = Document()
setup_page(doc2, style='journal')

add_title(doc2, '远程工作可行性如何缓解空气污染的经济损害')
add_abstract_and_keywords(doc2,
    '本文构建了Dingel-Neiman远程工作可行性指数的中国城市版本，'
    '并检验其是否调节PM2.5对城市GDP的边际损害。'
    '基于2000-2020年287个地级市面板数据，城市-年份双固定效应模型发现...',
    '远程工作；PM2.5；交互效应；调节效应')

add_heading_1(doc2, '一、引言')
add_body(doc2, '（引言内容略）')

add_page_number(doc2, show_total=True)
out2 = os.path.join(tempfile.gettempdir(), 'demo_journal.docx')
doc2.save(out2)
print(f'演示二（journal）已保存: {out2}')


# ═══════════════════════════════════════════════════
# 演示三：docxtpl 模板填充
# ═══════════════════════════════════════════════════
tmpdir = tempfile.gettempdir()
demo_template = os.path.join(tmpdir, 'demo_template.docx')
doc3 = Document()
setup_page(doc3, style='thesis')
add_title(doc3, '{{ title }}')
add_body(doc3, '{{ body }}')
doc3.save(demo_template)

rendered = os.path.join(tmpdir, 'demo_rendered.docx')
render_template(demo_template,
    {'title': '模板填充示例', 'body': '这段文字通过docxtpl的Jinja2模板引擎填充。'},
    rendered)
print(f'演示三（docxtpl模板）已保存: {rendered}')

print('\n=== 全部演示完成 ===')
