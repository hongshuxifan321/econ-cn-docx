# -*- coding: utf-8 -*-
"""econ-cn-docx 排版流水线：read JSON → render docx.

   Usage:
       from cn_docx import render
       render("paper_content.json", output="论文.docx")
"""

import json, os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ═══════════════════════════════════════════════════════════════
# Chinese quote repair
# ═══════════════════════════════════════════════════════════════

def _fix_quotes(text):
    """ASCII straight quotes -> Chinese curly quotes (paired).

    Only converts when the text contains CJK characters; pure ASCII/English
    text is left untouched, avoiding false positives like
    'He said "hello"' where a space precedes the quote.
    """
    if not text:
        return text
    # Quick pre-scan: if no CJK, leave ASCII quotes alone
    if not any(ord(ch) > 127 for ch in text):
        return text
    result = []
    in_quote = False
    for i, ch in enumerate(text):
        if ch == '"':
            prev = text[i-1] if i > 0 else ' '
            nxt = text[i+1] if i+1 < len(text) else ' '
            if not in_quote:
                if ord(prev) > 127 or prev in '，。、；：！？）】」\n\r\t' or i == 0:
                    result.append('“')
                    in_quote = True
                    continue
            else:
                if ord(nxt) > 127 or nxt in '，。、；：！？）】」\n\r\t' or i+1 >= len(text):
                    result.append('”')
                    in_quote = False
                    continue
        result.append(ch)
    if in_quote:
        result.append('”')
    return ''.join(result)


# ═══════════════════════════════════════════════════════════════
# Chinese font helper
# ═══════════════════════════════════════════════════════════════

def set_cn_font(run, font_name="宋体", size=Pt(12), bold=False):
    """Set w:ascii=TNR, w:hAnsi/w:eastAsia=font_name so Chinese
       punctuation (U+201C/D etc.) renders in the correct font."""
    run.font.name = 'Times New Roman'
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = size
    run.font.bold = bold


# ═══════════════════════════════════════════════════════════════
# OMML equation engine (LaTeX → Word native math)
# ═══════════════════════════════════════════════════════════════

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
XML_NS  = 'http://www.w3.org/XML/1998/namespace'

# LaTeX command → Unicode (rendered directly in OMML)
_LATEX_UNICODE = {
    '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
    '\\varepsilon': 'ε', '\\zeta': 'ζ', '\\eta': 'η', '\\theta': 'θ',
    '\\iota': 'ι', '\\kappa': 'κ', '\\lambda': 'λ', '\\mu': 'μ',
    '\\nu': 'ν', '\\xi': 'ξ', '\\pi': 'π', '\\rho': 'ρ',
    '\\sigma': 'σ', '\\tau': 'τ', '\\phi': 'φ', '\\chi': 'χ',
    '\\psi': 'ψ', '\\omega': 'ω',
    '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ',
    '\\Pi': 'Π', '\\Sigma': 'Σ', '\\Phi': 'Φ', '\\Psi': 'Ψ', '\\Omega': 'Ω',
    '\\cdot': '·', '\\times': '×', '\\div': '÷', '\\pm': '±',
    '\\leq': '≤', '\\geq': '≥', '\\neq': '≠', '\\approx': '≈',
    '\\to': '→', '\\rightarrow': '→', '\\Rightarrow': '⇒',
    '\\mid': '∣', '\\sum': '∑', '\\prod': '∏', '\\int': '∫',
    '\\partial': '∂', '\\infty': '∞', '\\nabla': '∇',
    '\\forall': '∀', '\\exists': '∃', '\\in': '∈', '\\notin': '∉',
    '\\subset': '⊂', '\\subseteq': '⊆', '\\sim': '∼', '\\equiv': '≡',
    '\\ln': 'ln', '\\log': 'log', '\\exp': 'exp',
    '\\max': 'max', '\\min': 'min',
}
# Combining diacritics
_DIACRITICS = {
    '\\bar':  '\N{COMBINING MACRON}',
    '\\hat':  '\N{COMBINING CIRCUMFLEX ACCENT}',
    '\\tilde': '\N{COMBINING TILDE}',
    '\\dot':  '\N{COMBINING DOT ABOVE}',
}


def _m(tag):
    return etree.Element(f'{{{MATH_NS}}}{tag}')


def _m_run(parent, text):
    r = etree.SubElement(parent, f'{{{MATH_NS}}}r')
    t = etree.SubElement(r, f'{{{MATH_NS}}}t')
    t.set(f'{{{XML_NS}}}space', 'preserve')
    t.text = text
    return r


def _m_box(parent, text):
    """Wrap text in a math box (m:box) inside an m:e for frac/rad args."""
    box = _m('box')
    e = etree.SubElement(box, f'{{{MATH_NS}}}e')
    # Recursively parse the inner text
    _render_math_content(e, text)
    return box


def _m_rad(parent, radicand, deg=None):
    """m:rad — square root. If deg is given, it's \\sqrt[deg]{radicand}."""
    rad = etree.SubElement(parent, f'{{{MATH_NS}}}rad')
    if deg is not None:
        deg_elem = etree.SubElement(rad, f'{{{MATH_NS}}}deg')
        _m_run(deg_elem, str(deg))
    e = etree.SubElement(rad, f'{{{MATH_NS}}}e')
    _render_math_content(e, radicand)


def _m_frac(parent, num_text, den_text):
    """m:f — fraction numerator/denominator."""
    f = etree.SubElement(parent, f'{{{MATH_NS}}}f')
    # numerator
    num = etree.SubElement(f, f'{{{MATH_NS}}}num')
    _render_math_content(num, num_text)
    # denominator
    den = etree.SubElement(f, f'{{{MATH_NS}}}den')
    _render_math_content(den, den_text)


def _m_sub(parent, base_text, sub_text):
    sSub = etree.SubElement(parent, f'{{{MATH_NS}}}sSub')
    e = etree.SubElement(sSub, f'{{{MATH_NS}}}e')
    _m_run(e, base_text)
    sub = etree.SubElement(sSub, f'{{{MATH_NS}}}sub')
    _m_run(sub, sub_text)


def _m_sup(parent, base_text, sup_text):
    sSup = etree.SubElement(parent, f'{{{MATH_NS}}}sSup')
    e = etree.SubElement(sSup, f'{{{MATH_NS}}}e')
    _m_run(e, base_text)
    sup = etree.SubElement(sSup, f'{{{MATH_NS}}}sup')
    _m_run(sup, sup_text)


def _m_sub_sup(parent, base_text, sub_text, sup_text):
    sSubSup = etree.SubElement(parent, f'{{{MATH_NS}}}sSubSup')
    e = etree.SubElement(sSubSup, f'{{{MATH_NS}}}e')
    _m_run(e, base_text)
    sub = etree.SubElement(sSubSup, f'{{{MATH_NS}}}sub')
    _m_run(sub, sub_text)
    sup = etree.SubElement(sSubSup, f'{{{MATH_NS}}}sup')
    _m_run(sup, sup_text)


def _normalize_commands(s):
    """Replace JSON-safe //cmd with \\cmd to avoid backslash escaping issues.
       //frac{a}{b} → \\frac{a}{b}, //sqrt{x} → \\sqrt{x}."""
    for cmd in ['frac', 'sqrt', 'partial', 'bar', 'hat', 'tilde', 'dot',
                'sum', 'prod', 'int', 'infty', 'times', 'cdot', 'pm',
                'leq', 'geq', 'neq', 'approx', 'to', 'rightarrow', 'Rightarrow']:
        s = s.replace('//' + cmd, '\\' + cmd)
    return s


# Whitelist: every backslash-command this system can handle.
# Built from _LATEX_UNICODE + _DIACRITICS + structural commands.
_VALID_LATEX_COMMANDS = (
    set(_LATEX_UNICODE.keys())
    | set(_DIACRITICS.keys())
    | {'\\frac', '\\sqrt', '\\text'}
)

# Map JSON escape chars back to the letter that follows a backslash.
_CTRL_TO_LETTER = {
    '\x07': 'a',   # bell      → \\a... (\\alpha)
    '\x08': 'b',   # backspace → \\b... (\\beta, \\bar)
    '\x09': 't',   # tab       → \\t... (\\tau, \\times, \\text, \\tilde, \\theta)
    '\x0a': 'n',   # newline   → \\n... (\\neq, \\nu)
    '\x0c': 'f',   # form feed → \\f... (\\frac)
    '\x0d': 'r',   # CR        → \\r... (\\rho, \\rightarrow, \\Rightarrow)
}


def _sanitize_latex(s):
    """Fix JSON-escape artifacts in LaTeX — only when the result is a known
       LaTeX command.  JSON may turn \\beta into <backspace>eta; we detect
       that <backspace> + 'eta' rebuilds \\beta, which is in the whitelist,
       so we restore it.  Unrecognised control chars are left untouched.
    """
    result = []
    i = 0
    while i < len(s):
        ch = s[i]
        if ch in _CTRL_TO_LETTER:
            letter = _CTRL_TO_LETTER[ch]
            # Scan ahead for letters to form a candidate command
            j = i + 1
            while j < len(s) and s[j].isalpha():
                j += 1
            candidate = '\\' + letter + s[i + 1:j]
            if candidate in _VALID_LATEX_COMMANDS:
                result.append(candidate)
                i = j
                continue
        result.append(ch)
        i += 1
    r = ''.join(result)
    # Also handle json.dump writing \\uXXXX escapes
    r = re.sub(r'\\\\u([0-9a-fA-F]{4})',
               lambda m: chr(int(m.group(1), 16)), r)
    return r


def _latex_to_unicode(s):
    """LaTeX commands → Unicode for direct OMML rendering."""
    s = _normalize_commands(s)
    s = _sanitize_latex(s)
    for cmd, uni in _LATEX_UNICODE.items():
        s = s.replace(cmd, uni)
    for cmd, diac in _DIACRITICS.items():
        # \\bar{x} → x + combining macron
        s = re.sub(rf'{re.escape(cmd)}\{{(\w)\}}', rf'\1{diac}', s)
    s = re.sub(r'\\text\{([^}]+)\}', r'\1', s)
    return s


def _find_matching_brace(s, start):
    """Given s[start] == '{', return index of matching '}'."""
    depth = 0
    for i in range(start, len(s)):
        if s[i] == '{':
            depth += 1
        elif s[i] == '}':
            depth -= 1
            if depth == 0:
                return i
    return -1


def _render_math_content(parent, text):
    """Render LaTeX-like math text into an OMML parent element.
    Handles: word_{sub}, word^{sup}, word_{sub}^{sup},
             \\frac{a}{b}, \\sqrt{x}, \\sqrt[n]{x},
             LaTeX Unicode commands, combining diacritics.
    """
    s = _latex_to_unicode(text)
    i = 0
    while i < len(s):
        # ── \\frac{num}{den} ──
        if s[i:].startswith('\\frac{') and i+6 < len(s):
            brace_start = i + 5  # after \\frac
            num_end = _find_matching_brace(s, brace_start)
            if num_end >= 0 and num_end + 1 < len(s) and s[num_end + 1] == '{':
                den_start = num_end + 1
                den_end = _find_matching_brace(s, den_start)
                if den_end >= 0:
                    num_text = s[brace_start+1:num_end]
                    den_text = s[den_start+1:den_end]
                    _m_frac(parent, num_text, den_text)
                    i = den_end + 1
                    continue

        # ── \\sqrt[n]{x} or \\sqrt{x} ──
        if s[i:].startswith('\\sqrt'):
            j = i + 5
            deg = None
            if j < len(s) and s[j] == '[':
                bracket_end = s.find(']', j)
                if bracket_end >= 0:
                    deg = s[j+1:bracket_end]
                    j = bracket_end + 1
            if j < len(s) and s[j] == '{':
                rad_end = _find_matching_brace(s, j)
                if rad_end >= 0:
                    _m_rad(parent, s[j+1:rad_end], deg)
                    i = rad_end + 1
                    continue

        # ── word_{sub}^{sup} or word^{sup}_{sub} ──
        m = re.match(r'(\S+?)_\{([^}]+)\}\^\{([^}]+)\}', s[i:])
        if not m:
            m = re.match(r'(\S+?)\^\{([^}]+)\}_\{([^}]+)\}', s[i:])
        if m:
            # Determine sub/sup order: which brace comes first in the source
            src = s[i:i+len(m.group(0))]
            underscore_pos = src.index('_') if '_' in src else 999
            caret_pos = src.index('^') if '^' in src else 999
            if underscore_pos < caret_pos:
                _m_sub_sup(parent, m.group(1), m.group(2), m.group(3))
            else:
                _m_sub_sup(parent, m.group(1), m.group(3), m.group(2))
            i += len(m.group(0))
            continue

        # ── word_{sub} ──
        m = re.match(r'(\S+?)_\{([^}]+)\}', s[i:])
        if m:
            _m_sub(parent, m.group(1), m.group(2))
            i += len(m.group(0))
            continue

        # ── word^{sup} ──
        m = re.match(r'(\S+?)\^\{([^}]+)\}', s[i:])
        if m:
            _m_sup(parent, m.group(1), m.group(2))
            i += len(m.group(0))
            continue

        # ── single-char sub/sup: x_2, x^2 ──
        m = re.match(r'(\S+)_(\S)', s[i:])
        if m:
            _m_sub(parent, m.group(1), m.group(2))
            i += len(m.group(0))
            continue
        m = re.match(r'(\S+)\^(\S)', s[i:])
        if m:
            _m_sup(parent, m.group(1), m.group(2))
            i += len(m.group(0))
            continue

        # ── plain text run ──
        _m_run(parent, s[i])
        i += 1


def _add_display_equation(doc, latex_str):
    """Insert centred OMML display equation (double-click to edit)."""
    omath = _m('oMath')
    _render_math_content(omath, latex_str)
    omath_para = etree.Element(f'{{{MATH_NS}}}oMathPara')
    omath_para.append(omath)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p._p.append(omath_para)


def _add_inline_omath(paragraph, math_text):
    """Add inline OMML run inside a paragraph."""
    omath = _m('oMath')
    _render_math_content(omath, math_text)
    run = paragraph.add_run()
    run._r.append(omath)


# ═══════════════════════════════════════════════════════════════
# RTF table parser (esttab output → python-docx table data)
# ═══════════════════════════════════════════════════════════════

def _rtf_u2uni(match):
    n = int(match.group(1))
    if n < 0:
        n += 65536
    return chr(n)


def _clean_rtf_text(text):
    text = re.sub(r'\\u(-?\d{2,5})\?', _rtf_u2uni, text)
    text = re.sub(r'\{\\super\s+([^}]+)\}', r'\1', text)
    text = re.sub(r'\{\\[a-z]+\d*\s+([^}]*)\}', r'\1', text)
    text = text.replace('{', '').replace('}', '')
    text = re.sub(r'\\[a-z]+\-?\d*\s?', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def _parse_rtf_table(filepath):
    """Parse esttab RTF → (headers, data_rows, notes, title)."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # title
    title = ""
    m = re.search(r'\{\\pard\\keepn\\ql\s(.+?)\\par\}', content)
    if not m:
        m = re.search(r'\{\\pard\\qc\\.+?\\fs\d+\s(.+?)\\par\}', content)
    if m:
        title = _clean_rtf_text(m.group(1))

    # notes — find \pard\ql block containing 注 after last table row
    notes = ""
    last_row = content.rfind('\\row')
    if last_row > 0:
        for m in re.finditer(r'\{\\pard\\ql\\fs\d+\s(.+?)\\par\}', content[last_row:]):
            if '\\u27880' in m.group(1):  # 注
                notes = _clean_rtf_text(m.group(1))
                break

    # rows
    raw_rows = []
    BS = chr(92)
    segments = re.split('(' + re.escape(BS + 'trowd') + ')', content)
    for k in range(1, len(segments), 2):
        segment = segments[k] + (segments[k+1] if k+1 < len(segments) else '')
        row_end = segment.find(BS + 'row')
        if row_end < 0:
            continue
        row_text = segment[:row_end]
        ncols = len(re.findall(re.escape(BS + 'cellx'), row_text))
        if ncols == 0:
            continue
        cell_parts = re.split(re.escape(BS + 'cell') + r'(?=\s|\\|$)', row_text)
        cells = []
        for part in cell_parts:
            part = part.strip()
            if not part:
                continue
            cm = re.search(re.escape(BS+'pard') + re.escape(BS+'intbl') +
                           re.escape(BS+'q') + r'[lrc]\s*(.+)$', part)
            cells.append(_clean_rtf_text(cm.group(1)) if cm else '')
        while len(cells) < ncols:
            cells.append('')
        if cells:
            raw_rows.append(cells[:ncols])

    if not raw_rows:
        return None, [], notes, title

    ncols = max(len(r) for r in raw_rows)

    # detect header
    headers = []
    data_start = 0
    for i, row in enumerate(raw_rows):
        if len(row) >= 2 and row[0]:
            headers = raw_rows[:i]
            data_start = i
            break
    if not headers:
        headers = raw_rows[:1]
        data_start = 1

    # merge multi-row headers
    if len(headers) >= 2:
        merged = []
        for j in range(ncols):
            parts = [h[j] for h in headers if j < len(h) and h[j]]
            merged.append('\n'.join(parts) if parts else '')
        header_row = merged
    elif len(headers) == 1:
        header_row = list(headers[0]) + [''] * (ncols - len(headers[0]))
    else:
        header_row = [''] * ncols

    # merge SE rows (empty first cell → standard errors)
    merged_rows = []
    for i in range(data_start, len(raw_rows)):
        row = list(raw_rows[i]) + [''] * (ncols - len(raw_rows[i]))
        if all(c == '' for c in row):
            continue
        if row[0] == '' and merged_rows:
            prev = merged_rows[-1]
            for j in range(1, ncols):
                if j < len(row) and row[j]:
                    se = row[j]
                    if not se.startswith('('):
                        se = '(' + se + ')'
                    prev[j] = prev[j] + '\n' + se if prev[j] else se
            continue
        merged_rows.append(row)

    return [header_row], merged_rows, notes, title


# ═══════════════════════════════════════════════════════════════
# Table rendering (RTF → Word 三线表)
# ═══════════════════════════════════════════════════════════════

def _set_cell_font(cell, text, font_name="宋体", size=Pt(9), bold=False,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with font; newlines → separate paragraphs."""
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    lines = text.split('\n') if text else ['']
    for k, line in enumerate(lines):
        if k > 0:
            p = cell.add_paragraph()
            p.alignment = alignment
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        run = p.add_run(_fix_quotes(line))
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        run.font.size = size
        run.font.bold = bold


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        if attrs:
            el = OxmlElement(f'w:{edge}')
            for key, val in attrs.items():
                el.set(qn(f'w:{key}'), str(val))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_rtf_table(doc, filepath, font_size=Pt(9)):
    """Parse esttab RTF and insert as 三线表."""
    headers, rows, notes, title = _parse_rtf_table(filepath)
    if headers is None:
        return

    # table title from RTF
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(_fix_quotes(title))
        set_cn_font(r, "黑体", Pt(10.5), bold=True)

    all_rows = headers + rows
    ncols = len(headers[0])
    nrows = len(all_rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row_data in enumerate(all_rows):
        for j, val in enumerate(row_data):
            cell = table.cell(i, j)
            is_header = (i == 0)
            _set_cell_font(cell, val if val else "",
                           font_name="黑体" if is_header else "宋体",
                           size=font_size,
                           bold=is_header,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 三线 borders
    for j in range(ncols):
        _set_cell_border(table.cell(0, j),
                         top={"sz": "12", "val": "single", "color": "000000"},
                         bottom={"sz": "8", "val": "single", "color": "000000"})
    for j in range(ncols):
        _set_cell_border(table.cell(nrows - 1, j),
                         bottom={"sz": "12", "val": "single", "color": "000000"})

    if notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(notes)
        set_cn_font(r, "宋体", Pt(9))


# ═══════════════════════════════════════════════════════════════
# Body paragraph (with inline math detection)
# ═══════════════════════════════════════════════════════════════

_MATH_PATTERN = re.compile(
    r'([A-Za-zα-ωΑ-Ωβγδε'
    r'λμσρπφθ∂ΣΠ'
    r'ΔΩ]+(?:_\{[^}]+\}|\^\{[^}]+\})+(?:_\{[^}]+\}|\^\{[^}]+\})*)'
)


def _add_body_paragraph(doc, text, size=Pt(12), indent=Pt(24),
                        line_spacing=1.5):
    """Body paragraph: justified, first-line indent, inline math detection."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = indent
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    parts = _MATH_PATTERN.split(text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 0:
            r = p.add_run(_fix_quotes(part))
            set_cn_font(r, "宋体", size)
        else:
            _add_inline_omath(p, part)


# ═══════════════════════════════════════════════════════════════
# Main render loop
# ═══════════════════════════════════════════════════════════════

def render(json_path, output="论文.docx", **kwargs):
    """Render paper_content.json → formatted .docx.

       Convenience wrapper: reads JSON then calls render_from_dicts().
       For programmatic use, call render_from_dicts(sections, output, ...) directly.
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        sections = json.load(f)
    render_from_dicts(sections, output, **kwargs)


def render_from_dicts(sections, output="论文.docx", *,
                      tables_dir="tables", figures_dir="figures",
                      table_map=None,
                      # page
                      margin_top=2.54, margin_bottom=2.54,
                      margin_left=3.18, margin_right=2.54,
                      # body
                      body_size=12, body_line_spacing=1.5, body_indent=24,
                      # headings
                      heading1_size=16, heading2_size=14, heading3_size=10.5,
                      # ref
                      ref_size=10, ref_line_spacing=1.25,
                      # table
                      table_font_size=9,
                      ):
    """Render a list of section dicts → formatted .docx.

       Each section dict: {"title": str, "body": str, "level": int}
       Same format as paper_content.json, but passed as a Python list.
    """
    doc = Document()

    # page setup
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(margin_top)
        section.bottom_margin = Cm(margin_bottom)
        section.left_margin = Cm(margin_left)
        section.right_margin = Cm(margin_right)

    # default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(body_size)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = body_line_spacing

    # resolve table paths
    if table_map is None:
        table_map = {}

    def _resolve_table(key):
        if key in table_map:
            return os.path.join(tables_dir, table_map[key])
        return os.path.join(tables_dir, key + '.rtf')

    def _resolve_image(name):
        return os.path.join(figures_dir, name)

    # ---- loop ----
    for sec in sections:
        title = _fix_quotes(sec.get("title", ""))
        body = _fix_quotes(sec.get("body", ""))
        level = sec.get("level", 0)

        # --- title block (level 0 + has title + not abstract/keywords) ---
        if level == 0 and title and title not in ('关键词', 'Keywords') \
                and not title.startswith('摘'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(title)
            set_cn_font(r, "黑体", Pt(18), bold=True)
            if body:
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.paragraph_format.space_after = Pt(12)
                r2 = p2.add_run(body)
                set_cn_font(r2, "黑体", Pt(14))
                body = ""

        # --- equation ---
        elif body.startswith('EQUATION:'):
            eq = body.replace('EQUATION:', '').strip()
            _add_display_equation(doc, eq)
            continue

        # --- abstract ---
        elif title == '摘要':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(title + '：')
            set_cn_font(r, "黑体", Pt(10.5), bold=True)
            if body:
                r2 = p.add_run(body)
                set_cn_font(r2, "楷体", Pt(10.5))
                body = ""

        # --- keywords ---
        elif title in ('关键词', 'Keywords'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(title + '：')
            set_cn_font(r, "黑体", Pt(10.5), bold=True)
            if body:
                r2 = p.add_run(body)
                set_cn_font(r2, "楷体", Pt(10.5))
                body = ""

        # --- L1 heading ---
        elif level == 1 and title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(title)
            set_cn_font(r, "黑体", Pt(heading1_size), bold=True)

        # --- L2 heading ---
        elif level == 2 and title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(title)
            set_cn_font(r, "黑体", Pt(heading2_size), bold=True)

        # --- L3 heading (figure/table caption) ---
        elif level == 3 and title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(title)
            set_cn_font(r, "黑体", Pt(heading3_size), bold=True)

        # ---- body text dispatch ----
        if not body:
            continue
        body = body.strip()
        if not body:
            continue

        if body.startswith("TABLE:"):
            key = body.replace("TABLE:", "").strip()
            fp = _resolve_table(key)
            if os.path.exists(fp):
                _add_rtf_table(doc, fp, font_size=Pt(table_font_size))
            else:
                print(f"[cn_docx] WARNING: table file not found: {fp}")
            continue

        if body.startswith("IMAGE:"):
            name = body.replace("IMAGE:", "").strip()
            fp = _resolve_image(name)
            if os.path.exists(fp):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(fp, width=Cm(10))
            else:
                print(f"[cn_docx] WARNING: image file not found: {fp}")
            continue

        if body.startswith("KEYWORDS:"):
            continue

        if body.startswith("REF:"):
            ref_text = body.replace("REF:", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = ref_line_spacing
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(ref_text)
            set_cn_font(r, "宋体", Pt(ref_size))
            continue

        if body.startswith("NOTE:"):
            p = doc.add_paragraph()
            r = p.add_run(body.replace("NOTE:", "").strip())
            set_cn_font(r, "宋体", Pt(9))
            r.italic = True
            continue

        # regular text — split on \n\n
        for para_text in body.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue

            if para_text.startswith("EQUATION:"):
                eq = para_text.replace("EQUATION:", "").strip()
                _add_display_equation(doc, eq)
                continue

            # H1/H2/H3 hypothesis line → bold
            if any(para_text.startswith(f"H{i}") for i in range(1, 5)) or \
               any(para_text.startswith(f"H{i}{c}")
                   for i in range(1, 5) for c in "ab"):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(body_indent)
                p.paragraph_format.line_spacing = body_line_spacing
                r = p.add_run(para_text)
                set_cn_font(r, "宋体", Pt(body_size), bold=True)
            else:
                _add_body_paragraph(doc, para_text,
                                    size=Pt(body_size),
                                    indent=Pt(body_indent),
                                    line_spacing=body_line_spacing)

    doc.save(output)
    if not output.startswith("__"):  # suppress print during tests
        print(f"论文已保存: {output}")
