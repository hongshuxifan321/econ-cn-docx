# -*- coding: utf-8 -*-
'econ-cn-docx 工具函数集：OMML 公式引擎 + 中文字体 + 引号修复 + 三线表 + 多风格预设'
import re, io
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree

# ═══════════════════════════════════════════════════
# 0. 排版风格预设
# ═══════════════════════════════════════════════════

STYLE_PRESETS = {
    'thesis': {  # 毕业论文 — GB/T 7713.1 + 多校经管类格式
        'margin_top': Cm(2.54), 'margin_bottom': Cm(2.54),
        'margin_left': Cm(3.18), 'margin_right': Cm(2.54),
        'body_size': Pt(12), 'body_line_spacing': 1.5,
        'body_indent': Pt(24),
        'heading1_size': Pt(16), 'heading2_size': Pt(14),
        'heading1_align': WD_ALIGN_PARAGRAPH.CENTER,
        'heading2_align': WD_ALIGN_PARAGRAPH.LEFT,
        'ref_size': Pt(10), 'ref_line_spacing': 1.25,
        'table_size': Pt(9),
    },
    'journal': {  # 期刊投稿 — 综合《经济研究》《经济学（季刊）》等
        'margin_top': Cm(2.54), 'margin_bottom': Cm(2.54),
        'margin_left': Cm(2.54), 'margin_right': Cm(2.54),
        'body_size': Pt(10.5), 'body_line_spacing': 1.5,
        'body_indent': Pt(21),
        'heading1_size': Pt(14), 'heading2_size': Pt(12),
        'heading1_align': WD_ALIGN_PARAGRAPH.CENTER,
        'heading2_align': WD_ALIGN_PARAGRAPH.LEFT,
        'ref_size': Pt(9), 'ref_line_spacing': 1.25,
        'table_size': Pt(8),
    },
    'working_paper': {  # 工作论文 — NBER/ElegantPaper 风格
        'margin_top': Cm(2.54), 'margin_bottom': Cm(2.54),
        'margin_left': Cm(2.54), 'margin_right': Cm(2.54),
        'body_size': Pt(11), 'body_line_spacing': 1.5,
        'body_indent': Pt(22),
        'heading1_size': Pt(14), 'heading2_size': Pt(12),
        'heading1_align': WD_ALIGN_PARAGRAPH.LEFT,
        'heading2_align': WD_ALIGN_PARAGRAPH.LEFT,
        'ref_size': Pt(10), 'ref_line_spacing': 1.25,
        'table_size': Pt(9),
    },
}

_current_style = STYLE_PRESETS['thesis']  # 默认毕业论文风格

def set_style(name):
    '切换排版风格：thesis / journal / working_paper'
    global _current_style
    if isinstance(name, str) and name in STYLE_PRESETS:
        _current_style = STYLE_PRESETS[name]
    elif isinstance(name, dict):
        _current_style = name

# ═══════════════════════════════════════════════════
# 1. 中文字体设置
# ═══════════════════════════════════════════════════

def set_cn_font(run, font_name='宋体', size=Pt(12), bold=False):
    '同时设置 w:ascii=TNR, w:hAnsi/w:eastAsia=中文字体。关键：中文引号（U+201C/D）属 hAnsi 范围'
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run.font.size = size
    run.bold = bold

# ═══════════════════════════════════════════════════
# 2. 中文弯引号修复
# ═══════════════════════════════════════════════════

def fix_chinese_quotes(text):
    '将 "含中文内容" 的 ASCII 直引号（U+0022）转为中文弯引号（U+201C/U+201D）'
    return re.sub(r'"([^"]*[一-鿿][^"]*)"', r'“\1”', text)

# ═══════════════════════════════════════════════════
# 3. OMML 公式引擎（LaTeX → Word 原生数学方程）
# ═══════════════════════════════════════════════════

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
XML_NS  = 'http://www.w3.org/XML/1998/namespace'

LATEX_UNICODE = {
    '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ',
    '\\delta': 'δ', '\\varepsilon': 'ε', '\\zeta': 'ζ',
    '\\eta': 'η', '\\theta': 'θ', '\\iota': 'ι',
    '\\kappa': 'κ', '\\lambda': 'λ', '\\mu': 'μ',
    '\\nu': 'ν', '\\xi': 'ξ', '\\pi': 'π',
    '\\rho': 'ρ', '\\sigma': 'σ', '\\tau': 'τ',
    '\\phi': 'φ', '\\chi': 'χ', '\\psi': 'ψ',
    '\\omega': 'ω',
    '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ',
    '\\Lambda': 'Λ', '\\Xi': 'Ξ', '\\Pi': 'Π',
    '\\Sigma': 'Σ', '\\Phi': 'Φ', '\\Psi': 'Ψ',
    '\\Omega': 'Ω',
    '\\cdot': '·', '\\times': '×', '\\div': '÷',
    '\\pm': '±', '\\leq': '≤', '\\geq': '≥',
    '\\neq': '≠', '\\approx': '≈',
    '\\rightarrow': '→', '\\Rightarrow': '⇒',
    '\\mid': '∣', '\\sum': '∑', '\\prod': '∏',
    '\\int': '∫',
    '\\ln': 'ln', '\\log': 'log', '\\exp': 'exp',
    '\\max': 'max', '\\min': 'min',
}

def _m(tag):
    return etree.Element(f'{{{MATH_NS}}}{tag}')

def _m_run(parent, text):
    r = etree.SubElement(parent, f'{{{MATH_NS}}}r')
    t = etree.SubElement(r, f'{{{MATH_NS}}}t')
    t.set(f'{{{XML_NS}}}space', 'preserve')
    t.text = text
    return r

def _m_sub(parent, base_text, sub_text):
    sSub = etree.SubElement(parent, f'{{{MATH_NS}}}sSub')
    e_elem = etree.SubElement(sSub, f'{{{MATH_NS}}}e')
    _m_run(e_elem, base_text)
    sub_elem = etree.SubElement(sSub, f'{{{MATH_NS}}}sub')
    _m_run(sub_elem, sub_text)
    return sSub

def _m_sup(parent, base_text, sup_text):
    sSup = etree.SubElement(parent, f'{{{MATH_NS}}}sSup')
    e_elem = etree.SubElement(sSup, f'{{{MATH_NS}}}e')
    _m_run(e_elem, base_text)
    sup_elem = etree.SubElement(sSup, f'{{{MATH_NS}}}sup')
    _m_run(sup_elem, sup_text)
    return sSup

def _m_sub_sup(parent, base_text, sub_text, sup_text):
    sSubSup = etree.SubElement(parent, f'{{{MATH_NS}}}sSubSup')
    e_elem = etree.SubElement(sSubSup, f'{{{MATH_NS}}}e')
    _m_run(e_elem, base_text)
    sub_elem = etree.SubElement(sSubSup, f'{{{MATH_NS}}}sub')
    _m_run(sub_elem, sub_text)
    sup_elem = etree.SubElement(sSubSup, f'{{{MATH_NS}}}sup')
    _m_run(sup_elem, sup_text)
    return sSubSup

def _latex_to_unicode(s):
    'LaTeX 命令 → Unicode，OMML 可直接渲染'
    for cmd, uni in LATEX_UNICODE.items():
        s = s.replace(cmd, uni)
    s = re.sub(r'\\text\{([^}]+)\}', r'\1', s)
    s = re.sub(r'\\tilde\{(\w)\}', r'\1̃', s)   # combining tilde
    s = re.sub(r'\\bar\{(\w)\}', r'\1̄', s)     # combining macron
    s = re.sub(r'\\hat\{(\w)\}', r'\1̂', s)     # combining circumflex
    return s

def _parse_latex_tokens(latex_str):
    '解析 LaTeX 为 token 流：下标、上标、上下标组合'
    tokens = []
    combo = r'(\S+?)_\{([^}]+)\}\^\{([^}]+)\}|(\S+?)\^\{([^}]+)\}_\{([^}]+)\}'
    braced = r'(\S+?)_\{([^}]+)\}|(\S+?)\^\{([^}]+)\}'
    single = r'(\S+?)_(\S)|(\S+?)\^(\S)'
    pattern = f'{combo}|{braced}|{single}'

    last_end = 0
    for m in re.finditer(pattern, latex_str):
        if m.start() > last_end:
            tokens.append(('text', latex_str[last_end:m.start()]))
        g = m.groups()
        if g[0] is not None:  # sub + sup
            tokens.append(('subsup', g[0], g[1], g[2]))
        elif g[3] is not None:  # sup + sub (reverse)
            tokens.append(('subsup', g[3], g[5], g[4]))
        elif g[6] is not None:  # sub only
            tokens.append(('sub', g[6], g[7]))
        elif g[8] is not None:  # sup only
            tokens.append(('sup', g[8], g[9]))
        elif g[10] is not None:  # single-char sub
            tokens.append(('sub', g[10], g[11]))
        elif g[12] is not None:  # single-char sup
            tokens.append(('sup', g[12], g[13]))
        last_end = m.end()
    if last_end < len(latex_str):
        tokens.append(('text', latex_str[last_end:]))
    return tokens

def _build_omath(tokens):
    '从 token 流构建 m:oMath 元素'
    omath = etree.Element(f'{{{MATH_NS}}}oMath')
    for token in tokens:
        if token[0] == 'text':
            _m_run(omath, token[1])
        elif token[0] == 'sub':
            _m_sub(omath, token[1], token[2])
        elif token[0] == 'sup':
            _m_sup(omath, token[1], token[2])
        elif token[0] == 'subsup':
            _m_sub_sup(omath, token[1], token[2], token[3])
    return omath

def add_display_equation(doc, latex_str):
    '插入 Word 原生显示方程（m:oMathPara），居中，双击可编辑'
    s = _latex_to_unicode(latex_str)
    tokens = _parse_latex_tokens(s)
    omath = _build_omath(tokens)
    omath_para = etree.Element(f'{{{MATH_NS}}}oMathPara')
    omath_para.append(omath)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p._p.append(omath_para)

def insert_inline_omath(paragraph, latex_str):
    '在段落当前位置插入 m:oMath 行内公式'
    s = _latex_to_unicode(latex_str)
    tokens = _parse_latex_tokens(s)
    omath = _build_omath(tokens)
    paragraph._p.append(omath)

def parse_inline_text(text, paragraph, base_size=Pt(12)):
    '解析行内 $...$ 公式（→ OMML）和 **粗体**，追加到段落'
    text = fix_chinese_quotes(text)
    tokens = re.split(r'(\$[^$]+\$)', text)
    for token in tokens:
        if token.startswith('$') and token.endswith('$'):
            insert_inline_omath(paragraph, token[1:-1])
        else:
            bold_tokens = re.split(r'(\*\*[^*]+\*\*)', token)
            for bt in bold_tokens:
                if bt.startswith('**') and bt.endswith('**'):
                    run = paragraph.add_run(bt[2:-2])
                    set_cn_font(run, '宋体', base_size, bold=True)
                else:
                    run = paragraph.add_run(bt)
                    set_cn_font(run, '宋体', base_size)

# ═══════════════════════════════════════════════════
# 4. 三线表
# ═══════════════════════════════════════════════════

def add_three_line_table(doc, headers, rows, col_widths=None):
    '添加三线表：顶线 12、表头下线 8、底线 12（单位 1/8 pt），字号由风格决定'
    s = _current_style
    ts = s['table_size']
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_text = fix_chinese_quotes(h)
        if '$' in h_text:
            for r in p.runs:
                r._element.getparent().remove(r._element)
            parse_inline_text(h_text, p, ts)
        else:
            run = p.add_run(h_text)
            set_cn_font(run, '黑体', ts, bold=True)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sval = fix_chinese_quotes(str(val))
            if '$' in sval:
                for r in p.runs:
                    r._element.getparent().remove(r._element)
                parse_inline_text(sval, p, ts)
            else:
                run = p.add_run(sval)
                set_cn_font(run, '宋体', ts)

    # 三线边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge, sz in [('top', 12), ('bottom', 12), ('insideH', 4)]:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

    # 表头下线加粗
    for j in range(len(headers)):
        cell = table.rows[0].cells[j]
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

    # 表头跨页重复
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    tblPr.append(tblHeader)

    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════
# 5. 页面设置
# ═══════════════════════════════════════════════════

def setup_page(doc, style='thesis'):
    'A4 + 根据风格预设设置边距和 Normal 样式基准'
    if isinstance(style, str):
        set_style(style)
    elif isinstance(style, dict):
        set_style(style)

    s = _current_style
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = s['margin_top']
        section.bottom_margin = s['margin_bottom']
        section.left_margin = s['margin_left']
        section.right_margin = s['margin_right']
    nstyle = doc.styles['Normal']
    nstyle.font.name = 'Times New Roman'
    nstyle.font.size = s['body_size']
    nstyle.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    nstyle.paragraph_format.line_spacing = s['body_line_spacing']

# ═══════════════════════════════════════════════════
# 6. 高级排版函数（直接调用，开箱即用）
# ═══════════════════════════════════════════════════

def _set_paragraph_spacing(paragraph, line_spacing=1.5, space_after=Pt(0)):
    '段落间距辅助函数'
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = Pt(0)

def add_title(doc, text):
    '论文标题：18pt 黑体居中加粗（所有风格统一）'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, line_spacing=1.5, space_after=Pt(12))
    run = p.add_run(fix_chinese_quotes(text))
    set_cn_font(run, '黑体', Pt(18), bold=True)

def add_heading_1(doc, text):
    '一级标题：字号和对齐方式由当前风格决定'
    s = _current_style
    p = doc.add_paragraph()
    p.alignment = s['heading1_align']
    _set_paragraph_spacing(p, line_spacing=s['body_line_spacing'], space_after=Pt(6))
    run = p.add_run(fix_chinese_quotes(text))
    set_cn_font(run, '黑体', s['heading1_size'], bold=True)

def add_heading_2(doc, text):
    '二级标题：字号和对齐方式由当前风格决定'
    s = _current_style
    p = doc.add_paragraph()
    p.alignment = s['heading2_align']
    _set_paragraph_spacing(p, line_spacing=s['body_line_spacing'], space_after=Pt(4))
    run = p.add_run(fix_chinese_quotes(text))
    set_cn_font(run, '黑体', s['heading2_size'], bold=True)

def add_heading_3(doc, text):
    '三级标题（表/图标题）：10.5pt 黑体居中加粗（所有风格统一）'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, line_spacing=1.5, space_after=Pt(2))
    run = p.add_run(fix_chinese_quotes(text))
    set_cn_font(run, '黑体', Pt(10.5), bold=True)

def add_body(doc, text, indent=True):
    '正文段落：字号、行距、缩进由当前风格决定，宋体两端对齐'
    s = _current_style
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p, line_spacing=s['body_line_spacing'])
    if indent:
        p.paragraph_format.first_line_indent = s['body_indent']
    parse_inline_text(text, p, s['body_size'])

def add_bullet(doc, text):
    '项目符号段落：字号、行距由当前风格决定，悬挂缩进'
    s = _current_style
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(p, line_spacing=s['body_line_spacing'])
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    parse_inline_text(text, p, s['body_size'])

def add_abstract(doc, body):
    '摘要：10.5pt 黑体加粗标签 + 楷体正文，同一段落左对齐顶格'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    _set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(0))
    r = p.add_run('摘要：')
    set_cn_font(r, '黑体', Pt(10.5), bold=True)
    r2 = p.add_run(fix_chinese_quotes(body))
    set_cn_font(r2, '楷体', Pt(10.5))

def add_keywords(doc, keywords):
    '关键词：10.5pt 黑体加粗标签 + 楷体内容，同一段落左对齐顶格'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    _set_paragraph_spacing(p, line_spacing=1.0, space_after=Pt(4))
    r = p.add_run('关键词：')
    set_cn_font(r, '黑体', Pt(10.5), bold=True)
    r2 = p.add_run(fix_chinese_quotes(keywords))
    set_cn_font(r2, '楷体', Pt(10.5))

def add_reference(doc, text):
    '参考文献条目：字号和行距由当前风格决定，宋体左对齐'
    s = _current_style
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, line_spacing=s['ref_line_spacing'])
    run = p.add_run(fix_chinese_quotes(text))
    set_cn_font(run, '宋体', s['ref_size'])

def add_table_note(doc, text):
    '表注：字号由当前风格决定，宋体左对齐'
    s = _current_style
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p, line_spacing=1.25)
    run = p.add_run(fix_chinese_quotes(text))
    set_cn_font(run, '宋体', s['table_size'])

# ═══════════════════════════════════════════════════
# 7. 扩展功能：模板、图片、页码、编号、图表、回归表
# ═══════════════════════════════════════════════════

# --- 7.1 模板加载 ---

def load_template(template_path):
    '从预设 docx 模板创建文档，自动继承模板的样式、多级编号、页边距'
    from docx import Document as Doc
    doc = Doc(template_path)
    set_style('thesis')  # 默认风格，用户可后续 setup_page 覆盖
    return doc

# --- 7.2 图片嵌入 ---

def add_figure(doc, image, caption='', width=Inches(5.5)):
    '插入图片（文件路径或 BytesIO），居中，图题在下方'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if isinstance(image, str):
        run.add_picture(image, width=width)
    else:
        run.add_picture(image, width=width)
    if caption:
        add_heading_3(doc, fix_chinese_quotes(caption))

# --- 7.3 页码 ---

def _insert_field(doc, run, field_code):
    '向 run 中插入 OOXML 域代码'
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_code} '
    run._r.append(instrText)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

def add_page_number(doc, show_total=True):
    '在页脚添加居中页码，默认格式 "第 X 页 / 共 Y 页"'
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('第 ')
        set_cn_font(run, '宋体', Pt(9))
        run2 = p.add_run()
        _insert_field(doc, run2, 'PAGE')
        set_cn_font(run2, '宋体', Pt(9))
        run3 = p.add_run(' 页')
        set_cn_font(run3, '宋体', Pt(9))
        if show_total:
            run4 = p.add_run(' / 共 ')
            set_cn_font(run4, '宋体', Pt(9))
            run5 = p.add_run()
            _insert_field(doc, run5, 'NUMPAGES')
            set_cn_font(run5, '宋体', Pt(9))
            run6 = p.add_run(' 页')
            set_cn_font(run6, '宋体', Pt(9))

# --- 7.4 多级标题自动编号 ---

def _get_counters(doc):
    '获取文档绑定的标题计数器（避免多文档串号）'
    if not hasattr(doc, '_heading_counters'):
        doc._heading_counters = {1: 0, 2: 0, 3: 0, 4: 0}
    return doc._heading_counters

def add_numbered_heading(doc, text, level=1):
    '带自动编号的标题：1, 1.1, 1.1.1...（计数器挂 doc，多文档安全）'
    c = _get_counters(doc)
    c[level] += 1
    for l in range(level + 1, 5):
        c[l] = 0
    prefix = '.'.join(str(c[l]) for l in range(1, level + 1))
    full = f'{prefix} {text}'
    if level == 1:
        add_heading_1(doc, full)
    elif level == 2:
        add_heading_2(doc, full)
    else:
        add_heading_3(doc, full)

def reset_heading_counters(doc):
    '重置标题计数器（开始新章节前调用）'
    c = _get_counters(doc)
    for k in c:
        c[k] = 0

# --- 7.5 matplotlib 图表嵌入 ---

def add_chart(doc, fig, caption='', width=Inches(5.5)):
    '将 matplotlib Figure 转为 PNG 嵌入 Word，居中，图题在下方'
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    add_figure(doc, buf, caption=caption, width=width)
    buf.close()

# --- 7.6 模板填充（docxtpl） ---

def render_template(template_path, context, output_path):
    '用 docxtpl 的 Jinja2 模板引擎填充 Word 模板并保存'
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    tpl.save(output_path)

# --- 7.7 封面 ---

def add_cover_page(doc, title, school='', author='', student_id='',
                   supervisor='', date='', department='', level='thesis'):
    '生成毕业论文封面：标题/学校/作者等居中排版'
    # 空行撑开顶部
    for _ in range(6):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, line_spacing=1.5)

    # 学校名
    if school:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(fix_chinese_quotes(school))
        set_cn_font(run, '黑体', Pt(22), bold=True)

    # 学位论文类型
    type_map = {'thesis': '硕士学位论文', 'bachelor': '本科学位论文', 'doctor': '博士学位论文'}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, line_spacing=1.5, space_after=Pt(24))
    run = p.add_run(type_map.get(level, '学位论文'))
    set_cn_font(run, '黑体', Pt(18), bold=True)

    # 论文标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, line_spacing=1.5, space_after=Pt(36))
    run = p.add_run(fix_chinese_quotes(title))
    set_cn_font(run, '黑体', Pt(22), bold=True)

    # 信息栏
    info_lines = []
    if author: info_lines.append(('姓    名', author))
    if student_id: info_lines.append(('学    号', student_id))
    if supervisor: info_lines.append(('指导教师', supervisor))
    if department: info_lines.append(('培养单位', department))
    if date: info_lines.append(('完成日期', date))

    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, line_spacing=1.5)
        run1 = p.add_run(f'{label}：')
        set_cn_font(run1, '宋体', Pt(14), bold=True)
        run2 = p.add_run(fix_chinese_quotes(value))
        set_cn_font(run2, '宋体', Pt(14))

    # 封面后分页
    doc.add_page_break()

# --- 7.8 摘要+关键词一键组合 ---

def add_abstract_and_keywords(doc, abstract_body, keywords):
    '摘要 + 关键词一键生成，中间自动空行'
    add_abstract(doc, abstract_body)
    add_keywords(doc, keywords)

# --- 7.9 回归结果专用表 ---

def add_regression_table(doc, models, var_names=None, notes='',
                         stars=True, digits=4):
    '根据回归结果 dict 列表生成标准回归报告表格'
    # models: [{'name': 'Model 1', 'N': 5000, 'R2': 0.123,
    #            'coef': {'x': 0.123, 'z': -0.045},
    #            'se':   {'x': 0.023, 'z': 0.018}}]
    if var_names is None:
        var_names = list(models[0]['coef'].keys())
    # 表头
    headers = ['变量'] + [m.get('name', f'({i+1})') for i, m in enumerate(models)]
    # 数据行
    rows = []
    for vn in var_names:
        row = [vn]
        for m in models:
            b = m['coef'].get(vn)
            se = m['se'].get(vn) if 'se' in m else None
            if b is None:
                row.append('')
            else:
                b_str = f'{b:.{digits}f}'
                if stars:
                    # 从 p 值推断显著性星号
                    p = m.get('pvalues', {}).get(vn, 1)
                    if p is not None:
                        if p < 0.01: b_str += '***'
                        elif p < 0.05: b_str += '**'
                        elif p < 0.1: b_str += '*'
                if se is not None:
                    row.append(f'{b_str}\n({se:.{digits}f})')
                else:
                    row.append(b_str)
        rows.append(row)
    # 附加统计
    add_three_line_table(doc, headers, rows)
    if notes:
        add_table_note(doc, notes)
    # 输出观测值和 R²
    stats_parts = []
    for m in models:
        n = m.get('N', '')
        r2 = m.get('R2', '')
        stats_parts.append(f"{m.get('name','')}: N={n}, $R^2$={r2:.4f}" if r2 else '')
    if any(stats_parts):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run('; '.join(stats_parts))
        set_cn_font(run, '宋体', _current_style['table_size'])
